"""
Document processing utilities for the OSINT tool.
Handles reading documents, extracting URLs, and saving output.
"""
import os
from docx import Document
from datetime import datetime

class DocumentProcessor:
    """Handles operations related to document processing."""
    
    @staticmethod
    def read_template(template_path):
        """
        Read the content of a template from a .docx file.
        
        Args:
            template_path: Path to the template file
            
        Returns:
            The template content as a string
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        doc = Document(template_path)
        template_content = ""
        for paragraph in doc.paragraphs:
            template_content += paragraph.text + "\n"
        
        return template_content.strip()
    
    @staticmethod
    def extract_urls(docx_path):
        """
        Extract URLs from a .docx file.
        
        Args:
            docx_path: Path to the document containing URLs
            
        Returns:
            List of URLs extracted from the document
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Document file not found: {docx_path}")
            
        doc = Document(docx_path)
        extracted_urls = []
        
        for paragraph in doc.paragraphs:
            if "http" in paragraph.text:
                url = paragraph.text.split("http", 1)[1]
                url = "http" + url.split()[0]
                extracted_urls.append(url.strip())
                
        return extracted_urls
    
    @staticmethod
    def save_report(content, input_file_name, output_dir="output_files"):
        """
        Save the generated report to a DOCX file.
        
        Args:
            content: The report content
            input_file_name: Original input file name (used for naming)
            output_dir: Directory to save the output
            
        Returns:
            Path to the saved file
        """
        # Create timestamp for unique filename
        datetime_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        input_file_base = os.path.splitext(os.path.basename(input_file_name))[0]
        output_file_name = f"{input_file_base}_{datetime_str}.docx"
        output_file_path = os.path.join(output_dir, output_file_name)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_file_path), exist_ok=True)
        
        # Create and save document
        doc = Document()
        doc.add_heading("Generated Report", level=1)
        doc.add_paragraph(content)
        doc.save(output_file_path)
        
        print(f"Report saved to {output_file_path}")
        return output_file_path
